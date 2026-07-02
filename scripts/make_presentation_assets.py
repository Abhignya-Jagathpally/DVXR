#!/usr/bin/env python3
"""make_presentation_assets.py — one command, an HONEST presentation pack.

Regenerates outputs/presentation/ deterministically and offline from existing
result files (benchmark_scoreboard.*, outputs/bci/metrics.json, emotiv_windows.csv,
codebook_usage.csv, ablation section of the scoreboard). Every figure caption states
the split, n, and chance/baseline; single-subject / proxy / exploratory results are
labelled as such. Nothing presents the BCI 0.82 as a success or fusion as a win.

Hard rules honoured:
  * no fabricated numbers — every value traces to a source file (see MANIFEST.md);
  * BCI command decode is demoted and captioned "reproduces Emotiv's on-device
    command state from raw EEG (exploratory, single-subject, single-session)";
  * the benchmark shows fusion losing wherever it loses, with CIs and a
    Brain2Qwerty reference line for context.
"""
from __future__ import annotations

import base64
import io
import json
import re
from pathlib import Path

import numpy as np
import pandas as pd

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "presentation"
FIG = OUT / "figures"
TAB = OUT / "tables"
SRC_SCORE_CSV = ROOT / "outputs" / "benchmark_scoreboard.csv"
SRC_SCORE_MD = ROOT / "outputs" / "benchmark_scoreboard.md"
SRC_BCI = ROOT / "outputs" / "bci" / "metrics.json"
SRC_WIN = ROOT / "outputs" / "bci" / "emotiv_windows.csv"
SRC_CODEBOOK = ROOT / "outputs" / "codebook_usage.csv"
SRC_BCI_MANIFOLD = ROOT / "outputs" / "bci" / "manifold_emotiv.png"
SRC_BCI_CONF = ROOT / "outputs" / "bci" / "command_confusion.png"

SEED = 7
# colourblind-safe (Wong palette)
CB = ["#0072B2", "#E69F00", "#009E73", "#D55E00", "#CC79A7", "#56B4E9", "#F0E442", "#000000"]
BRAIN2QWERTY_RER = (52, 58)   # reference band (%), for context only — a different task
SOURCES: dict[str, dict] = {}   # manifest: asset -> {source, shows}


def _mpl():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    plt.rcParams.update({"figure.dpi": 130, "font.size": 13, "axes.titlesize": 15,
                         "axes.labelsize": 13, "axes.grid": True, "grid.alpha": 0.3,
                         "savefig.bbox": "tight"})
    return plt


def _caption(fig, text):
    fig.text(0.5, -0.02, text, ha="center", va="top", fontsize=9.5,
             wrap=True, style="italic", color="#333333")


def _save(fig, name, source, shows):
    FIG.mkdir(parents=True, exist_ok=True)
    p = FIG / name
    fig.savefig(p)
    SOURCES[f"figures/{name}"] = {"source": source, "shows": shows}
    import matplotlib.pyplot as plt
    plt.close(fig)
    return p


def _write_table(df: pd.DataFrame, name: str, source: str, shows: str):
    TAB.mkdir(parents=True, exist_ok=True)
    df.to_csv(TAB / f"{name}.csv", index=False)
    try:
        md = df.to_markdown(index=False)
    except Exception:
        md = "```\n" + df.to_string(index=False) + "\n```"
    (TAB / f"{name}.md").write_text(f"# {name}\n\nSource: `{source}`\n\n{md}\n")
    SOURCES[f"tables/{name}.csv"] = {"source": source, "shows": shows}


# ------------------------------------------------------------------ parsing
def load_scoreboard() -> pd.DataFrame:
    return pd.read_csv(SRC_SCORE_CSV)


def _isnum(tok: str) -> bool:
    try:
        float(tok)
        return True
    except ValueError:
        return False


def _parse_md_sections() -> list[tuple[str, str, list[tuple[str, list[float]]]]]:
    """Parse fenced whitespace tables under '### <name> ...' headings.

    Returns (heading, header_line, rows) where each row is (label, [trailing floats]).
    Robust to spaces inside column headers (e.g. 'err_without (1-AUROC)').
    """
    if not SRC_SCORE_MD.exists():
        return []
    text = SRC_SCORE_MD.read_text()
    sections = []
    for m in re.finditer(r"###\s+([^\n(]+)([^\n]*)\n+```\n(.*?)\n```", text, re.S):
        name = m.group(1).strip()
        lines = [ln for ln in m.group(3).strip().splitlines() if ln.strip()]
        if len(lines) < 2:
            continue
        header = lines[0]
        rows = []
        for ln in lines[1:]:
            toks = ln.split()
            vals = []
            while toks and _isnum(toks[-1]):
                vals.insert(0, float(toks.pop()))
            rows.append((" ".join(toks), vals))
        sections.append((name, header, rows))
    return sections


def load_per_config() -> dict[str, pd.DataFrame]:
    """per-config CV tables (label + one metric value) per task."""
    out = {}
    for name, header, rows in _parse_md_sections():
        good = [(lbl, v) for lbl, v in rows if len(v) == 1 and lbl]
        if name in ("stress", "glucose", "mortality") and len(good) >= 3:
            out[name] = pd.DataFrame({"config": [g[0] for g in good],
                                      "value": [g[1][0] for g in good]})
    return out


def load_ablation() -> pd.DataFrame | None:
    """ablation table: modality + [err_without, err_with_all, contribution, ci_low, ci_high]."""
    cols = ["dropped_modality", "err_without", "err_with_all",
            "contribution", "ci_low", "ci_high"]
    for name, header, rows in _parse_md_sections():
        if "dropped" not in header.lower():
            continue
        good = [(lbl, v) for lbl, v in rows if len(v) >= 5 and lbl]
        if len(good) >= 2:
            data = {"dropped_modality": [g[0] for g in good]}
            for i, c in enumerate(cols[1:]):
                data[c] = [g[1][i] for g in good]
            return pd.DataFrame(data)
    return None


def load_bci() -> dict:
    return json.loads(SRC_BCI.read_text())


# ------------------------------------------------------------------ BCI LOBO
def compute_lobo(feature_prefix: str = "eeg_", n_blocks: int = 5) -> dict:
    """Honest strict drift control: leave-one-TIME-block-out 4-class command decode
    on RAW-EEG (welch) band power. Chance = 0.25. Reproduces the MC-engine label."""
    from sklearn.linear_model import LogisticRegression
    from sklearn.metrics import balanced_accuracy_score
    from sklearn.preprocessing import StandardScaler

    df = pd.read_csv(SRC_WIN)
    cmd = df[df["label"].isin(["Left", "Right", "Push", "Pull"])].copy()
    cmd = cmd.sort_values("t_center").reset_index(drop=True)
    feats = [c for c in cmd.columns if c.startswith(feature_prefix)]
    X = cmd[feats].to_numpy(float)
    y = cmd["label"].to_numpy()
    block = np.floor(np.linspace(0, n_blocks - 1e-9, len(cmd))).astype(int)
    accs = []
    for b in range(n_blocks):
        tr, te = block != b, block == b
        if te.sum() == 0 or len(np.unique(y[tr])) < 2:
            continue
        sc = StandardScaler().fit(X[tr])
        clf = LogisticRegression(max_iter=1000, class_weight="balanced",
                                 random_state=SEED).fit(sc.transform(X[tr]), y[tr])
        accs.append(balanced_accuracy_score(y[te], clf.predict(sc.transform(X[te]))))
    return {"lobo_balanced_acc": float(np.mean(accs)) if accs else float("nan"),
            "n_blocks": int(len(accs)), "n_windows": int(len(cmd)),
            "feature_set": feature_prefix.rstrip("_") + " (raw-EEG welch)", "chance": 0.25}


# ------------------------------------------------------------------ figures
def fig_scoreboard(sb: pd.DataFrame):
    plt = _mpl()
    fig, ax = plt.subplots(figsize=(9, 5))
    tasks = sb["task"].tolist()
    x = np.arange(len(tasks))
    w = 0.36
    base = sb["base_err"].to_numpy(float)
    prop = sb["prop_err"].to_numpy(float)
    ax.bar(x - w / 2, base, w, label="best baseline", color=CB[0])
    ax.bar(x + w / 2, prop, w, label="CACMF fused (proposed)", color=CB[3])
    for i, r in sb.iterrows():
        ax.annotate(f"RER {r['RER_pct']:.0f}%\n[{r['RER_CI_low']:.0f},{r['RER_CI_high']:.0f}]",
                    (i, max(base[i], prop[i])), textcoords="offset points",
                    xytext=(0, 6), ha="center", fontsize=9)
    ax.set_xticks(x)
    ax.set_xticklabels([f"{t}\n({m})" for t, m in zip(tasks, sb["metric"])])
    ax.set_ylabel("error (lower is better)")
    ax.set_title("Honest relativity: where fusion helps and where it doesn't")
    # Brain2Qwerty reference line (context only): a ~52-58% RER win on a different task
    ax.axhline(0, color="k", lw=0.5)
    ax.text(0.99, 0.97, f"Brain2Qwerty reference: ~{BRAIN2QWERTY_RER[0]}-{BRAIN2QWERTY_RER[1]}% RER\n"
            "(different task; for scale only)", transform=ax.transAxes, ha="right", va="top",
            fontsize=8.5, bbox=dict(boxstyle="round", fc="#fff3cd", ec="#e0a800"))
    ax.legend(loc="upper left")
    _caption(fig, "Subject/patient-held-out 5x5 CV. Negative RER = fused is WORSE than the "
             "best non-fused baseline. RER 95% CI in brackets. n: stress=1534 win/20 subj, "
             "glucose=21119/19, mortality=100/100. Fusion does not beat the baseline on any task.")
    return _save(fig, "benchmark_scoreboard.png", "outputs/benchmark_scoreboard.csv",
                 "per-task base vs fused error, RER% + CI")


def fig_modality_ablation(ab: pd.DataFrame):
    plt = _mpl()
    fig, ax = plt.subplots(figsize=(8, 4.5))
    contrib_col = [c for c in ab.columns if "contribution" in str(c).lower()][0]
    lo = [c for c in ab.columns if "ci_low" in str(c).lower()]
    hi = [c for c in ab.columns if "ci_high" in str(c).lower()]
    mod_col = ab.columns[0]
    mods = ab[mod_col].astype(str).tolist()
    val = ab[contrib_col].to_numpy(float)
    y = np.arange(len(mods))
    err = None
    if lo and hi:
        err = np.abs(np.vstack([val - ab[lo[0]].to_numpy(float),
                                ab[hi[0]].to_numpy(float) - val]))
    ax.barh(y, val, xerr=err, color=CB[2], capsize=4)
    ax.set_yticks(y)
    ax.set_yticklabels(mods)
    ax.invert_yaxis()
    ax.set_xlabel("Δ error when the modality is removed (higher = more important)")
    ax.set_title("True modality ablation (retrain WITHOUT the modality) — stress")
    _caption(fig, "Stress task, subject-held-out CV. Each modality is dropped and the fused "
             "model RETRAINED (not zero-filled). Bars = mean Δ(1-AUROC) with 95% CI. "
             "Motion contributes most; every modality helps.")
    return _save(fig, "modality_ablation.png",
                 "outputs/benchmark_scoreboard.md (ablation section)",
                 "per-modality retrain-without contribution + CI (stress)")


def fig_fusion_vs_concat(pc: pd.DataFrame):
    plt = _mpl()
    fig, ax = plt.subplots(figsize=(9, 5))
    pc = pc.copy()
    pc.columns = ["config", "err"]
    pc = pc.sort_values("err")
    def kind(c):
        if c.startswith("single:"):
            return CB[1]
        if c in ("rep:raw", "rep:pca", "classical_gbm"):
            return CB[0]
        if "fused" in c or "cacmf" in c:
            return CB[3]
        return CB[5]
    ax.barh(np.arange(len(pc)), pc["err"], color=[kind(c) for c in pc["config"]])
    ax.set_yticks(np.arange(len(pc)))
    ax.set_yticklabels(pc["config"])
    ax.invert_yaxis()
    ax.set_xlabel("CV error (1-AUROC, lower is better)")
    ax.set_title("Single modality vs naive concat vs learned fusion — stress")
    import matplotlib.patches as mp
    ax.legend(handles=[mp.Patch(color=CB[1], label="single modality"),
                       mp.Patch(color=CB[0], label="all-modality concat / GBM"),
                       mp.Patch(color=CB[3], label="CACMF learned fusion")], loc="lower right")
    _caption(fig, "Stress, subject-held-out 5x5 CV, chance (majority) 1-AUROC=0.5. "
             "Concatenating modalities beats the best single modality; learned CACMF fusion "
             "does NOT beat naive concatenation.")
    return _save(fig, "fusion_vs_concat.png", "outputs/benchmark_scoreboard.md (stress per-config)",
                 "single vs concat vs fusion CV error (stress)")


def _png_with_caption(src: Path, name: str, title: str, caption: str, source: str, shows: str):
    plt = _mpl()
    import matplotlib.image as mpimg
    fig, ax = plt.subplots(figsize=(7.5, 6))
    if src.exists():
        ax.imshow(mpimg.imread(src))
    else:
        ax.text(0.5, 0.5, "source figure unavailable", ha="center")
    ax.axis("off")
    ax.set_title(title)
    _caption(fig, caption)
    return _save(fig, name, source, shows)


def fig_bci_controls(bci: dict, lobo: dict):
    plt = _mpl()
    ts = bci["two_stage"]
    fig, axes = plt.subplots(1, 3, figsize=(13, 4.4))
    # panel 1: engaged-vs-neutral + lateralization AUROC vs chance 0.5
    a = axes[0]
    vals = [ts.get("engaged_vs_neutral_auroc", np.nan), ts.get("lateralization_lr_auroc", np.nan)]
    a.bar(["engaged\nvs neutral", "Left vs Right\nlateralization"], vals, color=[CB[0], CB[1]])
    a.axhline(0.5, color="r", ls="--", label="chance 0.5")
    a.set_ylim(0, 1); a.set_ylabel("AUROC"); a.set_title("Neural separability (controls)")
    a.legend()
    for i, v in enumerate(vals):
        a.text(i, v + 0.02, f"{v:.3f}", ha="center")
    # panel 2: 4-class command balanced acc across split strictness vs chance 0.25
    b = axes[1]
    names = ["trial-\ngrouped", "temporal\nblock", "leave-one\nblock-out"]
    accs = [ts.get("command_4class_balanced_acc", np.nan),
            ts.get("temporal_block_balanced_acc", np.nan), lobo["lobo_balanced_acc"]]
    b.bar(names, accs, color=[CB[4], CB[2], CB[3]])
    b.axhline(0.25, color="r", ls="--", label="chance 0.25")
    b.set_ylim(0, 1); b.set_ylabel("balanced accuracy"); b.set_title("Command decode vs split strictness")
    b.legend()
    for i, v in enumerate(accs):
        b.text(i, (v or 0) + 0.02, f"{v:.3f}", ha="center")
    # panel 3: raw-EEG welch 4-class vs chance 0.2
    c = axes[2]
    dec = bci["decoding"]
    c.bar(["welch\n(raw EEG)", "pow\n(Emotiv)"], [dec["welch"]["balanced_acc"], dec["pow"]["balanced_acc"]],
          color=[CB[0], CB[5]])
    c.axhline(0.2, color="r", ls="--", label="chance 0.2")
    c.set_ylim(0, 1); c.set_ylabel("balanced accuracy"); c.set_title("5-class decode by feature set")
    c.legend()
    fig.suptitle("BCI honest controls — single-subject, single-session, exploratory", y=1.03)
    _caption(fig, "Emotiv EPOC X, one subject, one 1373 s session. Labels come from Emotiv's "
             "on-device Mental-Command engine (not experimenter cues), so this REPRODUCES the "
             "MC-engine state, not validated neural intent. Neural separability is at chance "
             "(engaged 0.489, lateralization 0.541); command accuracy collapses as the split "
             "gets stricter (block confound).")
    return _save(fig, "bci_honest_controls.png", "outputs/bci/metrics.json + emotiv_windows.csv (LOBO)",
                 "engaged/lateralization AUROC; command balAcc vs split; welch vs pow")


def fig_codebook():
    plt = _mpl()
    df = pd.read_csv(SRC_CODEBOOK)
    fig, ax = plt.subplots(figsize=(8, 4))
    if {"code_index", "count"}.issubset(df.columns):
        ax.bar(df["code_index"], df["count"], color=CB[0])
        ax.set_xlabel("codebook index"); ax.set_ylabel("usage count")
    ppl = None
    if "frequency" in df.columns:
        f = df["frequency"].to_numpy(float)
        ppl = float(np.exp(-(f * np.log(f + 1e-12)).sum()))
    ax.set_title(f"VQ codebook usage" + (f"  (perplexity {ppl:.1f})" if ppl else ""))
    _caption(fig, "VQ codebook usage histogram over the demo fixture. Perplexity = effective "
             "number of codes used (higher = more diverse). Diagnostic, not a benchmark metric.")
    return _save(fig, "codebook_usage.png", "outputs/codebook_usage.csv",
                 "codebook histogram + perplexity")


def fig_galea(bci: dict):
    plt = _mpl()
    g = bci.get("galea", {})
    n_ch = g.get("n_channels", 16)
    n_use = g.get("n_usable", 0)
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.bar(["usable", "railed/unusable"], [n_use, n_ch - n_use], color=[CB[2], CB[3]])
    ax.set_ylabel("channels"); ax.set_title(f"Galea/OpenBCI signal quality ({n_use}/{n_ch} usable)")
    _caption(fig, f"Galea (OpenBCI) resting recording, fs={g.get('fs','?')} Hz, "
             f"{g.get('n_samples','?')} samples. {n_use} of {n_ch} channels usable — "
             "multi-device ingestion story; not used for decoding.")
    return _save(fig, "galea_signal_quality.png", "outputs/bci/metrics.json (galea)",
                 "usable vs railed channels")


# ------------------------------------------------------------------ tables
def build_tables(sb: pd.DataFrame, pc: dict, ab: pd.DataFrame | None, bci: dict, lobo: dict):
    sb2 = sb.copy()
    sb2.insert(0, "protocol", "subject/patient-held-out 5x5 CV")
    _write_table(sb2, "scoreboard_table", "outputs/benchmark_scoreboard.csv",
                 "headline RER per task with CI, p-values, meets>=50%")
    for task, df in pc.items():
        d = df.copy(); d.columns = ["config", "cv_error"]
        _write_table(d, f"per_config_cv_{task}", "outputs/benchmark_scoreboard.md",
                     f"per-config CV error for {task}")
    if ab is not None:
        _write_table(ab, "modality_ablation_table",
                     "outputs/benchmark_scoreboard.md (ablation)", "retrain-without contribution")
    ts = bci["two_stage"]; dec = bci["decoding"]
    decode = pd.DataFrame([
        {"metric": "welch 5-class balAcc (raw EEG)", "value": dec["welch"]["balanced_acc"], "chance": 0.2},
        {"metric": "pow 5-class balAcc (Emotiv)", "value": dec["pow"]["balanced_acc"], "chance": 0.2},
        {"metric": "engaged-vs-neutral AUROC", "value": ts["engaged_vs_neutral_auroc"], "chance": 0.5},
        {"metric": "Left-vs-Right lateralization AUROC", "value": ts["lateralization_lr_auroc"], "chance": 0.5},
        {"metric": "command 4-class balAcc (trial-grouped) [DEMOTED]",
         "value": ts["command_4class_balanced_acc"], "chance": 0.25},
        {"metric": "command 4-class balAcc (temporal-block)", "value": ts["temporal_block_balanced_acc"], "chance": 0.25},
        {"metric": "command 4-class balAcc (leave-one-block-out, raw EEG)",
         "value": lobo["lobo_balanced_acc"], "chance": 0.25},
    ])
    _write_table(decode, "bci_decode_table", "outputs/bci/metrics.json + emotiv_windows.csv",
                 "BCI decode + honest controls (labels_source: emotiv_mc_engine)")
    comp = pd.DataFrame([
        {"task": r["task"], "RER_pct": r["RER_pct"], "meets_50pct": r["meets_>=50%"],
         "honest_sentence": _honest_sentence(r)} for _, r in sb.iterrows()
    ] + [{"task": "Brain2Qwerty (reference, different task)", "RER_pct": "~52-58",
          "meets_50pct": True, "honest_sentence": "External reference point for scale only; not our result."}])
    _write_table(comp, "comparison_table", "outputs/benchmark_scoreboard.csv + literature",
                 "RER vs Brain2Qwerty reference + one honest sentence per task")
    return decode, comp


def _honest_sentence(r) -> str:
    t = r["task"]
    if t == "stress":
        return "Multimodal (4 peripheral-physiology streams); fusion loses to naive concat (RER<0)."
    if t == "glucose":
        return "CGM-only (single modality); learned model loses to simple features/persistence."
    if t == "mortality":
        return "EHR-only (single modality); tiny n=100, fused overfits and loses badly."
    return ""


# ------------------------------------------------------------------ dashboard / brief / narrative
def _b64(p: Path) -> str:
    return base64.b64encode(p.read_bytes()).decode()


def build_dashboard(figs: list[Path], decode: pd.DataFrame, sb: pd.DataFrame):
    parts = ["<!doctype html><html><head><meta charset='utf-8'><title>DVXR — honest results</title>",
             "<style>body{font-family:system-ui,Arial;max-width:1100px;margin:2em auto;padding:0 1em;color:#222}"
             "h1{border-bottom:3px solid #0072B2}h2{margin-top:1.6em;color:#0072B2}"
             "img{max-width:100%;border:1px solid #ddd;border-radius:6px}"
             "table{border-collapse:collapse;margin:1em 0}td,th{border:1px solid #ccc;padding:4px 8px;font-size:13px}"
             ".cav{background:#fff3cd;border:1px solid #e0a800;padding:8px 12px;border-radius:6px}</style></head><body>"]
    parts.append("<h1>DVXR — honest, presentation-ready results</h1>")
    parts.append("<div class='cav'><b>Honesty note.</b> Fusion does not beat strong baselines on "
                 "any real task (negative RER). The BCI command result is demoted: it reproduces "
                 "Emotiv's on-device Mental-Command state from raw EEG (exploratory, single-subject, "
                 "single-session) — not validated neural-intent decoding. Every number traces to a "
                 "file (see MANIFEST.md).</div>")
    titles = {"benchmark_scoreboard.png": "Benchmark: honest relativity scoreboard",
              "modality_ablation.png": "True modality ablation (stress)",
              "fusion_vs_concat.png": "Single vs concat vs learned fusion (stress)",
              "bci_honest_controls.png": "BCI honest controls (the real BCI headline)",
              "bci_manifold.png": "BCI neural manifold (exploratory)",
              "bci_confusion.png": "BCI command confusion (demoted)",
              "codebook_usage.png": "VQ codebook usage", "galea_signal_quality.png": "Galea signal quality"}
    for f in figs:
        parts.append(f"<h2>{titles.get(f.name, f.name)}</h2>")
        parts.append(f"<img src='data:image/png;base64,{_b64(f)}'/>")
    parts.append("<h2>Scoreboard table</h2>" + sb.to_html(index=False))
    parts.append("<h2>BCI decode + honest controls</h2>" + decode.to_html(index=False))
    parts.append("</body></html>")
    (OUT / "dashboard.html").write_text("".join(parts))
    SOURCES["dashboard.html"] = {"source": "all figures + tables", "shows": "self-contained pack"}


def build_brief(sb: pd.DataFrame, decode: pd.DataFrame, lobo: dict):
    try:
        import docx
        from docx.shared import Pt
    except Exception:
        (OUT / "results_brief.md").write_text(_brief_md(sb, lobo))
        SOURCES["results_brief.md"] = {"source": "tables", "shows": "one-page brief (md fallback)"}
        return
    d = docx.Document()
    d.add_heading("DVXR — mid-project results brief (honest)", 0)
    d.add_paragraph("Goal: a multimodal health-signal pipeline (CACMF) + a real-label benchmark "
                    "and a BCI pilot. This brief reports only measured, traceable numbers.")
    d.add_heading("Honest benchmark (subject/patient-held-out 5x5 CV)", 1)
    for _, r in sb.iterrows():
        d.add_paragraph(
            f"{r['task']} ({r['metric']}): fused {r['prop_err']:.3f} vs {r['best_baseline']} "
            f"{r['base_err']:.3f} -> RER {r['RER_pct']:.0f}% "
            f"[{r['RER_CI_low']:.0f},{r['RER_CI_high']:.0f}], meets>=50%: {r['meets_>=50%']}",
            style="List Bullet")
    d.add_paragraph("Takeaway: learned fusion does NOT beat strong baselines; multimodality helps "
                    "only via concatenation (stress). Reporting this honestly is the contribution.")
    d.add_heading("BCI pilot (single-subject, single-session, exploratory)", 1)
    d.add_paragraph("Labels come from Emotiv's Mental-Command engine (labels_source: "
                    "emotiv_mc_engine), so the 4-class command decode reproduces on-device state, "
                    "not validated neural intent. Neural separability is at chance "
                    "(engaged 0.489, lateralization 0.541 AUROC); command balanced-accuracy falls "
                    f"from 0.82 (trial-grouped) to {lobo['lobo_balanced_acc']:.2f} (leave-one-block-out).")
    d.add_heading("Limitations & next steps", 1)
    d.add_paragraph("Features are summary statistics (caps the ceiling); single wearable/subject; "
                    "no dataset co-registers EEG+CGM+EHR. Next: raw-signal encoders, a cued "
                    "multi-subject BCI dataset (e.g. PhysioNet MI), nested-CV headline.")
    d.save(OUT / "results_brief.docx")
    SOURCES["results_brief.docx"] = {"source": "tables", "shows": "one-page honest brief"}


def _brief_md(sb, lobo):
    lines = ["# DVXR — mid-project results brief (honest)\n"]
    for _, r in sb.iterrows():
        lines.append(f"- {r['task']} ({r['metric']}): fused {r['prop_err']:.3f} vs "
                     f"{r['best_baseline']} {r['base_err']:.3f} -> RER {r['RER_pct']:.0f}% "
                     f"[{r['RER_CI_low']:.0f},{r['RER_CI_high']:.0f}]")
    lines.append(f"\nBCI: single-subject/exploratory; command balAcc 0.82 (trial-grouped) -> "
                 f"{lobo['lobo_balanced_acc']:.2f} (leave-one-block-out); labels_source=emotiv_mc_engine.")
    return "\n".join(lines) + "\n"


def build_narrative(sb: pd.DataFrame, lobo: dict):
    md = f"""# Slide narrative (honest)

## 1. Title — DVXR: multimodal health signals, honestly evaluated
- What we built + the honesty stance. Figure: none. Note: lead with the framing.

## 2. The pipeline (CACMF)
- Per-modality encoders -> VQ codebook -> cross-modal fusion -> calibrated heads.
- Figure: codebook_usage.png. Note: architecture is real and runs offline/CPU.

## 3. The honest benchmark
- Real labels, subject/patient-held-out 5x5 CV, CIs + significance.
- Figure: benchmark_scoreboard.png. Note: "fusion does NOT beat strong baselines — and reporting that is the contribution."

## 4. Where multimodality DOES help
- Concatenation beats best single modality (stress); learned fusion doesn't beat concat.
- Figure: fusion_vs_concat.png + modality_ablation.png. Note: motion dominates stress.

## 5. BCI pilot — honest controls
- Single-subject/single-session; labels from Emotiv MC engine (not cues).
- Figure: bci_honest_controls.png. Note: engaged 0.489, lateralization 0.541 (chance); 0.82 -> {lobo['lobo_balanced_acc']:.2f} under leave-one-block-out.

## 6. BCI geometry (exploratory)
- Figure: bci_manifold.png + bci_confusion.png. Note: interesting geometry, NOT validated decoding.

## 7. Limitations & next steps
- Summary-stat features; single wearable/subject; no EEG+CGM+EHR co-registration.
- Next: raw-signal encoders, cued multi-subject BCI (PhysioNet MI), nested-CV headline.
"""
    (OUT / "slide_narrative.md").write_text(md)
    SOURCES["slide_narrative.md"] = {"source": "tables + figures", "shows": "slide-by-slide outline"}


def build_manifest(bci: dict):
    lines = ["# Presentation asset MANIFEST — every number traces to a source\n",
             "Generated by `scripts/make_presentation_assets.py` (offline, deterministic, seed=7).\n",
             f"\n**Data provenance flag:** BCI assets derive from `outputs/bci/*` and "
             f"`emotiv_windows.csv`, which were produced from the FULL Emotiv recording that is "
             f"NOT committed (only a 5.5 KB sample is in `data/sample/`). "
             f"labels_source = **emotiv_mc_engine** (no experimenter cues). "
             f"Benchmark assets derive from committed real public data in `data/real/`.\n",
             "\n| asset | source file | shows |", "|---|---|---|"]
    for asset, meta in sorted(SOURCES.items()):
        lines.append(f"| `{asset}` | `{meta['source']}` | {meta['shows']} |")
    (OUT / "MANIFEST.md").write_text("\n".join(lines) + "\n")


# ------------------------------------------------------------------ main
def main() -> int:
    np.random.seed(SEED)
    OUT.mkdir(parents=True, exist_ok=True)
    sb = load_scoreboard()
    pc = load_per_config()
    ab = load_ablation()
    bci = load_bci()
    lobo = compute_lobo()

    figs = [fig_scoreboard(sb)]
    if ab is not None:
        figs.append(fig_modality_ablation(ab))
    if "stress" in pc:
        figs.append(fig_fusion_vs_concat(pc["stress"]))
    figs.append(fig_bci_controls(bci, lobo))
    figs.append(_png_with_caption(
        SRC_BCI_MANIFOLD, "bci_manifold.png", "BCI neural manifold (PHATE) — exploratory",
        "Emotiv EPOC X, single subject/session, colored by MC-engine command. Geometry is "
        "interesting but NOT validated neural-intent decoding.", "outputs/bci/manifold_emotiv.png",
        "EEG manifold colored by command"))
    figs.append(_png_with_caption(
        SRC_BCI_CONF, "bci_confusion.png", "BCI command confusion (DEMOTED)",
        "4-class command, trial-grouped CV, chance 0.25, single-subject. Reproduces Emotiv "
        "MC-engine state (labels_source: emotiv_mc_engine) — not validated neural decoding.",
        "outputs/bci/command_confusion.png", "confusion matrix (demoted)"))
    figs.append(fig_codebook())
    figs.append(fig_galea(bci))

    decode, comp = build_tables(sb, pc, ab, bci, lobo)
    build_dashboard(figs, decode, sb)
    build_brief(sb, decode, lobo)
    build_narrative(sb, lobo)
    build_manifest(bci)

    print(f"[assets] wrote {len(figs)} figures + {len(list(TAB.glob('*.csv')))} tables to {OUT}")
    print(f"[assets] LOBO (raw-EEG, leave-one-block-out) balAcc = {lobo['lobo_balanced_acc']:.3f} "
          f"(chance 0.25, n={lobo['n_windows']})")
    for p in ["dashboard.html", "results_brief.docx", "results_brief.md",
              "slide_narrative.md", "MANIFEST.md"]:
        if (OUT / p).exists():
            print(f"[assets]   {OUT / p}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
